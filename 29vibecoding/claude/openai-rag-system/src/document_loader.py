"""
Document Loader for RAG system
.doc, .docx 파일을 로드하고 청킹하는 모듈
"""

from pathlib import Path
from typing import List
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from config import DATA_DIR, CHUNK_SIZE, CHUNK_OVERLAP, SUPPORTED_EXTENSIONS


class DocumentLoader:
    """문서 로딩 및 청킹 클래스"""
    
    def __init__(self, data_dir: Path = DATA_DIR):
        self.data_dir = data_dir
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def load_docx(self, file_path: Path) -> str:
        """
        .docx 파일을 로드하여 텍스트 추출
        
        Args:
            file_path: 문서 파일 경로
            
        Returns:
            추출된 텍스트
        """
        try:
            doc = DocxDocument(file_path)
            full_text = []
            
            # 모든 단락 추출
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # 표 내용 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        full_text.append(row_text)
            
            return "\n".join(full_text)
            
        except Exception as e:
            print(f"❌ 파일 로드 실패 {file_path.name}: {str(e)}")
            return ""
    
    def load_documents(self) -> List[Document]:
        """
        data 폴더의 모든 문서를 로드
        
        Returns:
            LangChain Document 객체 리스트
        """
        documents = []
        files = []
        
        # 지원하는 확장자 파일 찾기
        for ext in SUPPORTED_EXTENSIONS:
            files.extend(self.data_dir.glob(f"*{ext}"))
        
        if not files:
            raise FileNotFoundError(
                f"{self.data_dir}에 문서 파일이 없습니다. "
                f"지원 형식: {', '.join(SUPPORTED_EXTENSIONS)}"
            )
        
        print(f"\n📁 총 {len(files)}개 문서 발견")
        
        # 각 파일 로드
        for file_path in files:
            print(f"📄 로딩 중: {file_path.name}")
            
            if file_path.suffix in [".doc", ".docx"]:
                text = self.load_docx(file_path)
                
                if text:
                    doc = Document(
                        page_content=text,
                        metadata={
                            "source": file_path.name,
                            "file_path": str(file_path),
                            "file_type": file_path.suffix
                        }
                    )
                    documents.append(doc)
                    print(f"   ✅ 성공 (길이: {len(text)} 문자)")
                else:
                    print(f"   ⚠️  빈 문서")
        
        print(f"\n✅ 총 {len(documents)}개 문서 로드 완료\n")
        return documents
    
    def split_documents(self, documents: List[Document]) -> List[Document]:
        """
        문서를 청크로 분할
        
        Args:
            documents: 원본 문서 리스트
            
        Returns:
            분할된 Document 청크 리스트
        """
        print(f"✂️  문서 청킹 중... (크기: {CHUNK_SIZE}, 오버랩: {CHUNK_OVERLAP})")
        
        chunks = self.text_splitter.split_documents(documents)
        
        print(f"✅ 총 {len(chunks)}개 청크 생성 완료\n")
        
        return chunks
    
    def load_and_split(self) -> List[Document]:
        """
        문서 로드 및 청킹을 한번에 수행
        
        Returns:
            분할된 Document 청크 리스트
        """
        documents = self.load_documents()
        chunks = self.split_documents(documents)
        return chunks


def test_loader():
    """문서 로더 테스트"""
    print("=== Document Loader 테스트 ===\n")
    
    loader = DocumentLoader()
    
    try:
        chunks = loader.load_and_split()
        
        if chunks:
            print(f"\n--- 첫 번째 청크 미리보기 ---")
            print(f"소스: {chunks[0].metadata['source']}")
            print(f"내용: {chunks[0].page_content[:200]}...")
            print(f"전체 길이: {len(chunks[0].page_content)} 문자")
            
    except Exception as e:
        print(f"❌ 오류: {str(e)}")


if __name__ == "__main__":
    test_loader()
